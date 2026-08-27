#!/usr/bin/env python3
"""Build the editable, two-page Word version of Benjamin Berman's resume."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "BenjaminBermanSoftwareEngineerResume.docx"
FONT = "Arial"
ACCENT = RGBColor(31, 78, 121)


def set_cell_free_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)


def set_font(run, size: float, *, bold=False, italic=False, color=None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_repeatable_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    section = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    section.font.name = FONT
    section.font.size = Pt(11.5)
    section.font.bold = True
    section.font.color.rgb = ACCENT
    section._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    section.paragraph_format.space_before = Pt(5)
    section.paragraph_format.space_after = Pt(2)
    section.paragraph_format.keep_with_next = True

    job = styles.add_style("Resume Job", WD_STYLE_TYPE.PARAGRAPH)
    job.font.name = FONT
    job.font.size = Pt(10.35)
    job._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    job.paragraph_format.space_before = Pt(3)
    job.paragraph_format.space_after = Pt(1)
    job.paragraph_format.keep_with_next = True

    body = styles.add_style("Resume Body", WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = FONT
    body.font.size = Pt(10.2)
    body._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    body.paragraph_format.space_after = Pt(1.5)
    body.paragraph_format.line_spacing = 1.0

    bullet = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.font.name = FONT
    bullet.font.size = Pt(10.1)
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.13)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.line_spacing = 1.0

    sub_bullet = styles.add_style("Resume Sub-bullet", WD_STYLE_TYPE.PARAGRAPH)
    sub_bullet.font.name = FONT
    sub_bullet.font.size = Pt(9.9)
    sub_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    sub_bullet.paragraph_format.left_indent = Inches(0.4)
    sub_bullet.paragraph_format.first_line_indent = Inches(-0.13)
    sub_bullet.paragraph_format.space_after = Pt(0.8)
    sub_bullet.paragraph_format.line_spacing = 1.0


def bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    borders.append(bottom)
    p_pr.append(borders)


def add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Resume Section")
    paragraph.add_run(title.upper())
    bottom_border(paragraph)


def add_body(document: Document, text: str) -> None:
    document.add_paragraph(text, style="Resume Body")


def add_skill(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph(style="Resume Body")
    set_font(paragraph.add_run(f"{label}: "), 10.1, bold=True)
    set_font(paragraph.add_run(text), 10.1)
    paragraph.paragraph_format.space_after = Pt(0.8)


def add_job(
    document: Document,
    employer: str,
    location: str,
    role: str,
    dates: str,
    *,
    page_break_before: bool = False,
) -> None:
    paragraph = document.add_paragraph(style="Resume Job")
    paragraph.paragraph_format.page_break_before = page_break_before
    set_font(paragraph.add_run(employer), 10.35, bold=True)
    set_font(paragraph.add_run(f" — {location}"), 10.35)
    set_font(paragraph.add_run(f" | {role}"), 10.35, italic=True)
    set_font(paragraph.add_run(f" | {dates}"), 10.35)


def add_bullet(document: Document, text: str, *, level=0) -> None:
    style = "Resume Bullet" if level == 0 else "Resume Sub-bullet"
    paragraph = document.add_paragraph(style=style)
    marker = "•" if level == 0 else "◦"
    set_font(paragraph.add_run(f"{marker}  {text}"), 10.1 if level == 0 else 9.9)


def add_header(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run("Benjamin Berman"), 18, bold=True, color=ACCENT)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run("SOFTWARE ENGINEER & TECHNICAL LEAD"), 10, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_font(paragraph.add_run("Pittsburgh, PA  •  bnjmnbrmn@gmail.com"), 9)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    set_font(
        paragraph.add_run(
            "linkedin.com/in/bnjmnbrmn  •  bnjmnbrmn.com  •  github.com/bnjmnbrmn"
        ),
        9,
    )


def build(output: Path) -> None:
    document = Document()
    set_cell_free_page_layout(document)
    set_repeatable_styles(document)
    document.core_properties.title = "Benjamin Berman — Software Engineer & Technical Lead"
    document.core_properties.author = "Benjamin Berman"
    document.core_properties.subject = "Resume"
    document.core_properties.keywords = (
        "Java, Spring Boot, REST APIs, microservices, Oracle, technical leadership, AI-assisted development"
    )

    add_header(document)

    add_section(document, "Professional Summary")
    add_body(
        document,
        "I am a software engineer and technical lead with a Computer Science Ph.D. and 11+ years of "
        "industry experience, including financial systems at BNY and Interactive Brokers. Java and Spring "
        "Boot are my bread and butter, but my broader focus is making complex systems reliable, "
        "understandable, and easier to change. I am an enthusiastic, pragmatic adopter of AI-assisted "
        "development, using coding agents within a development process grounded in static typing, "
        "automated testing, and human review to accelerate delivery without giving up correctness. I also "
        "enjoy mentoring developers and creating tools and workflows that help both human and AI "
        "contributors work effectively.",
    )

    add_section(document, "Technical Skills")
    add_skill(
        document,
        "Core",
        "Java, Spring Boot, Spring MVC, Spring WebFlux, Spring Data JPA, microservices, REST, JSON, "
        "Java EE, Hibernate/JPA, Apache ActiveMQ Artemis, Kafka, Kafka Streams, Project Reactor",
    )
    add_skill(
        document,
        "Data",
        "SQL, Oracle/PL/SQL, PostgreSQL, MySQL, MongoDB, Redis, Elasticsearch, Flyway, Jackson",
    )
    add_skill(
        document,
        "Testing & Delivery",
        "AI-assisted development, end-to-end testing, JUnit, TestNG, Git, GitLab CI/CD, Gradle, Maven, "
        "Docker, Kubernetes, Terraform, Ansible, Linux/Unix, Bash/Zsh",
    )
    add_skill(
        document,
        "Cloud, Security & Observability",
        "AWS (CCP certified; EC2, EKS, MSK, S3, RDS, SQS, API Gateway, CloudWatch), OAuth, JWT, "
        "SSL/TLS, Kibana, Grafana, nginx, Tomcat, JBoss",
    )
    add_skill(
        document,
        "Additional",
        "Scala, Ruby, JavaScript, TypeScript, Angular, AngularJS, SOAP, XML, YAML, Ant, Mercurial, "
        "Subversion, Coq, Alloy, Lustre, ESC/Java2, Swing",
    )

    add_section(document, "Professional Experience")
    add_job(
        document,
        "eNGINE",
        "Pittsburgh, PA",
        "Senior Software Engineer — BNY engagement",
        "June 2025–Present",
    )
    add_bullet(
        document,
        "Extended an asynchronous integration between BNY's Spring Boot/Angular digital-assets platform "
        "and its enterprise information warehouse to support financial reporting; implemented and tested "
        "local messaging workflows with Apache ActiveMQ Artemis.",
    )
    add_bullet(
        document,
        "Flagged the blast-radius risk of enabling reporting across approximately 60 transaction types, "
        "partnered with Product to define the initial scope, built AI-assisted end-to-end test coverage, "
        "and implemented per-transaction-type configuration to support incremental rollout.",
        level=1,
    )
    add_bullet(
        document,
        "Adapted quickly to a major late-stage requirement for authoritative system-of-record identifiers, "
        "working with CIW and Product stakeholders to refocus reporting on completed, fee-bearing "
        "blockchain transactions.",
        level=1,
    )
    add_bullet(
        document,
        "Updated the Oracle database schema and indexes to associate off-chain platform records with "
        "corresponding on-chain activity.",
        level=1,
    )
    add_bullet(
        document,
        "Built an AI-assisted, multi-repository development environment with local service orchestration, "
        "end-to-end testing, parallel builds, and Bash/Zsh tab completion; optional local-database support "
        "reduced application startup from over two minutes to under 20 seconds. Presented the environment "
        "to other engineers and helped colleagues adopt it.",
    )

    add_job(
        document,
        "TeleTracking Technologies Inc",
        "Pittsburgh, PA",
        "Software Engineer III",
        "July 2024–March 2025",
    )
    add_bullet(
        document,
        "Debugged and fixed hospital logistics software on an event-oriented platform written in Java, C#, "
        "and Python.",
    )

    add_job(
        document,
        "UPMC Enterprises",
        "Pittsburgh, PA",
        "Senior Software Engineer",
        "February 2022–April 2024",
        page_break_before=True,
    )
    add_body(
        document,
        "Contributed to a large-scale medical-records aggregation and analysis platform spanning a wide "
        "range of microservices and web applications.",
    )
    add_bullet(document, "Used reactive streams to improve throughput by more than 3×.")
    add_bullet(document, "Debugged the CI/CD pipeline to improve some job completion times by up to 8×.")

    add_job(
        document,
        "M*Modal / 3M",
        "Pittsburgh, PA",
        "Senior Software Engineer; included service as tech lead",
        "July 2017–February 2022",
    )
    add_body(
        document,
        "Developed Java/Spring Boot services with architectural, leadership, and mentoring responsibilities "
        "for Computer Assisted Provider Documentation (CAPD), which connected a physician-facing interface "
        "to real-time natural-language processing.",
    )
    add_bullet(
        document,
        "Designed and implemented Kafka Streams microservices that processed and aggregated approximately "
        "1,000 requests per second in real time; coordinated across application, reporting, legacy ETL, "
        "product, and sales teams to resolve critical ambiguities.",
    )
    add_bullet(
        document,
        "Connected real-time NLP with prior document analysis, handling identifier mappings, configuration "
        "and credential caching, OAuth, fallback mechanisms, additional API calls, and observability.",
    )
    add_bullet(
        document,
        "Created reusable Terraform and Ansible automation for consistent deployment across AWS environments.",
    )

    add_job(
        document,
        "Interactive Brokers Group",
        "Greenwich, CT",
        "Secure Application Developer",
        "February 2015–July 2017",
    )
    add_bullet(
        document,
        "Developed secure financial applications using Java EE and Oracle/PL/SQL, with an emphasis on "
        "encryption and multi-factor authentication.",
    )

    add_job(
        document,
        "The University of Iowa",
        "Iowa City, IA",
        "Research/Teaching Assistant and Instructor",
        "August 2008–December 2014",
    )
    add_bullet(
        document,
        "Taught and developed material for introductory Java; supported undergraduate and graduate courses "
        "including HCI and Formal Methods.",
    )
    add_bullet(
        document,
        "Developed and evaluated user interfaces for programming tools and helped secure NSF grant "
        "CCF-1250306.",
    )

    add_section(document, "Education")
    add_body(document, "Ph.D., Computer Science — The University of Iowa, December 2014")
    add_body(document, "B.S., Physics — Rice University, May 2008")

    add_section(document, "Publications & Presentations")
    add_body(
        document,
        "Berman, B. A. and Hourcade, J. P. (2014). “Keyboard Card Menus: A New Presentation of "
        "Non-Standard Shortcuts.” Journal of Universal Computer Science.",
    )
    add_body(
        document,
        "Berman, B. A. and Hourcade, J. P. (2014). “Keyboard Card Menus: Faster Learning of Many Fast "
        "Commands.” XIV International Congress of Human-Computer Interaction (Interaction 2013).",
    )
    add_body(
        document,
        "“New Coq User Interfaces: Survey and Ideas.” Co-authored presentation, Coq Workshop 2012, "
        "Princeton, NJ.",
    )

    # Avoid compatibility mode and keep Word/LibreOffice pagination stable.
    settings = document.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    destination = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_OUTPUT
    build(destination)
    print(destination)
